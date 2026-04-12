# Copyright 2025 Softwell S.r.l. - Licensed under Apache License 2.0
# See LICENSE file for details

"""WordCompiler - Compiler for Word documents (.docx).

Transforms a built Bag (with WordBuilder) into a Word document using
the genro-builders @compiler dispatch and auto-walk. Handlers receive
the node and parent live object. Resolved attributes are accessed via
node.runtime_attrs, resolved value via node.runtime_value.
"""

from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from genro_bag import Bag
from genro_builders import BagCompilerBase
from genro_builders.compiler import compiler

if TYPE_CHECKING:
    from genro_bag import BagNode


class WordCompiler(BagCompilerBase):
    """Compiler Word: Bag -> bytes (.docx).

    Uses @compiler handlers with parent propagation.
    Maintains a live Document and _live_map for incremental updates.
    """

    def __init__(self, builder: Any) -> None:
        super().__init__(builder)
        self._doc: Any | None = None
        self._live_map: dict[int, Any] = {}
        self._custom_handlers: dict[str, Any] = {}

    def register_handler(self, tag: str, handler: Any) -> None:
        """Register a custom compile handler for a tag.

        The handler receives (node, parent) and is called during document
        building. Custom handlers take precedence over @compiler ones.

        Args:
            tag: The element tag name (e.g., "qrcode").
            handler: Callable(node, parent).
        """
        self._custom_handlers[tag] = handler

    def _dispatch_compile(self, node: BagNode, parent: Any = None) -> Any | None:
        """Override dispatch to check custom handlers first."""
        tag = node.node_tag or node.label

        custom = self._custom_handlers.get(tag)
        if custom is not None:
            return custom(node, parent)

        return super()._dispatch_compile(node, parent=parent)

    # -------------------------------------------------------------------------
    # Compile / Render / Serialize
    # -------------------------------------------------------------------------

    def compile(self, built_bag: Bag, target: Any = None) -> Any:  # noqa: ARG002
        """Compile built Bag into a live Document."""
        self._doc = Document()
        self._live_map.clear()
        list(self._walk_compile(built_bag, parent=self._doc))
        return self._doc

    def render(self, compiled_bag: Bag) -> bytes:
        """Compile and serialize to bytes (.docx)."""
        self.compile(compiled_bag)
        return self.serialize()

    def serialize(self) -> bytes:
        """Serialize the live Document to bytes without rebuilding."""
        if self._doc is None:
            return b""
        buffer = BytesIO()
        self._doc.save(buffer)
        return buffer.getvalue()

    # -------------------------------------------------------------------------
    # Live update
    # -------------------------------------------------------------------------

    def update_node(self, node: BagNode) -> bool:
        """Try to update the live object for a node."""
        live_obj = self._live_map.get(id(node))
        if live_obj is None:
            return False
        attrs = node.runtime_attrs
        return self._apply_live_update(attrs, node.node_tag or "", live_obj)

    def _apply_live_update(
        self, attrs: dict[str, Any], tag: str, live_obj: Any
    ) -> bool:
        """Apply resolved attribute changes to a live docx object."""
        if tag in ("run", "cell"):
            self._apply_run_formatting(attrs, live_obj)
            content = attrs.get("content", "")
            if content:
                live_obj.text = str(content)
            return True

        if tag in ("paragraph", "heading", "item"):
            if hasattr(live_obj, "runs") and live_obj.runs:
                content = attrs.get("content", "")
                if content and live_obj.runs:
                    live_obj.runs[0].text = str(content)
                self._apply_run_formatting(attrs, live_obj.runs[0])
            align = attrs.get("align")
            if align:
                live_obj.alignment = self._get_alignment(align)
            return True

        return False

    # -------------------------------------------------------------------------
    # @compiler handlers — auto-walked by _dispatch_compile
    # -------------------------------------------------------------------------

    @compiler()
    def document(self, node: BagNode, parent: Any) -> Any:
        """Compile document node. Parent is the Document object."""
        doc = parent
        attrs = node.runtime_attrs
        section = doc.sections[0]

        orientation = attrs.get("orientation")
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width = section.page_height
            new_height = section.page_width
            section.page_width = new_width
            section.page_height = new_height

        for margin_name in ("top", "bottom", "left", "right"):
            margin_val = attrs.get(f"margin_{margin_name}")
            if margin_val is not None:
                setattr(section, f"{margin_name}_margin", Cm(margin_val))

        title = attrs.get("title", "")
        if title:
            doc.add_heading(title, level=0)

        return doc

    @compiler()
    def heading(self, node: BagNode, parent: Any) -> Any:
        """Compile heading node. Parent is the Document."""
        attrs = node.runtime_attrs
        content = attrs.get("content", "")
        level = attrs.get("level", 1)

        hdg = parent.add_heading(str(content), level=level)
        self._live_map[id(node)] = hdg

        if hdg.runs:
            self._apply_run_formatting(attrs, hdg.runs[0])

        return hdg

    @compiler()
    def paragraph(self, node: BagNode, parent: Any) -> Any:
        """Compile paragraph node. Parent is Document, cell, header, or footer."""
        attrs = node.runtime_attrs
        content = attrs.get("content", "")
        style = attrs.get("style")

        if hasattr(parent, "add_paragraph"):
            para = parent.add_paragraph(style=style) if style else parent.add_paragraph()
        else:
            para = parent

        self._live_map[id(node)] = para

        if content:
            run = para.add_run(str(content))
            self._apply_run_formatting(attrs, run)

        align = attrs.get("align")
        if align:
            para.alignment = self._get_alignment(align)

        space_before = attrs.get("space_before")
        if space_before is not None:
            para.paragraph_format.space_before = Pt(space_before)

        space_after = attrs.get("space_after")
        if space_after is not None:
            para.paragraph_format.space_after = Pt(space_after)

        line_spacing = attrs.get("line_spacing")
        if line_spacing is not None:
            para.paragraph_format.line_spacing = line_spacing

        return para

    @compiler()
    def run(self, node: BagNode, parent: Any) -> Any:
        """Compile run node. Parent is a paragraph or heading."""
        attrs = node.runtime_attrs
        content = attrs.get("content", "")
        docx_run = parent.add_run(str(content))
        self._apply_run_formatting(attrs, docx_run)
        self._live_map[id(node)] = docx_run
        return docx_run

    @compiler()
    def itemlist(self, node: BagNode, parent: Any) -> Any:
        """Compile itemlist node. Stores list_type for child items."""
        attrs = node.runtime_attrs
        self._current_list_type = attrs.get("type", "bullet")
        return parent

    @compiler()
    def item(self, node: BagNode, parent: Any) -> Any:
        """Compile item node. Parent is the Document (passed through from itemlist)."""
        attrs = node.runtime_attrs
        content = attrs.get("content", "")
        list_type = getattr(self, "_current_list_type", "bullet")
        style = "List Number" if list_type == "number" else "List Bullet"
        para = parent.add_paragraph(str(content), style=style)
        self._live_map[id(node)] = para
        return para

    @compiler()
    def table(self, node: BagNode, parent: Any) -> Any:
        """Compile table node. Parent is the Document."""
        attrs = node.runtime_attrs
        num_cols = 0
        if isinstance(node.value, Bag):
            for child in node.value:
                if child.node_tag == "row" and isinstance(child.value, Bag):
                    num_cols = len(child.value)
                    break

        tbl = parent.add_table(rows=0, cols=num_cols)

        style = attrs.get("style")
        if style:
            tbl.style = style

        align = attrs.get("align")
        if align:
            align_map = {
                "left": WD_TABLE_ALIGNMENT.LEFT,
                "center": WD_TABLE_ALIGNMENT.CENTER,
                "right": WD_TABLE_ALIGNMENT.RIGHT,
            }
            if align.lower() in align_map:
                tbl.alignment = align_map[align.lower()]

        return tbl

    @compiler()
    def row(self, node: BagNode, parent: Any) -> Any:
        """Compile row node. Parent is the table."""
        attrs = node.runtime_attrs
        table_row = parent.add_row()

        row_height = attrs.get("height")
        if row_height:
            table_row.height = Cm(row_height)

        self._current_col_idx = 0
        return table_row

    @compiler()
    def cell(self, node: BagNode, parent: Any) -> Any:
        """Compile cell node. Parent is the table row."""
        attrs = node.runtime_attrs
        col_idx = getattr(self, "_current_col_idx", 0)
        docx_cell = parent.cells[col_idx]
        self._current_col_idx = col_idx + 1

        content = attrs.get("content", "")
        width = attrs.get("width")
        bold = attrs.get("bold", False)
        bg_color = attrs.get("bg_color")
        align = attrs.get("align")
        valign = attrs.get("valign")

        if width:
            docx_cell.width = Cm(width)

        if bg_color:
            self._set_cell_shading(docx_cell, bg_color)

        if valign:
            valign_map = {
                "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
                "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
            }
            if valign.lower() in valign_map:
                docx_cell.vertical_alignment = valign_map[valign.lower()]

        if content:
            para = docx_cell.paragraphs[0]
            para.clear()
            docx_run = para.add_run(str(content))
            self._live_map[id(node)] = docx_run

            if bold:
                docx_run.bold = True

            self._apply_run_formatting(attrs, docx_run)

            if align:
                para.alignment = self._get_alignment(align)

        return docx_cell

    @compiler()
    def image(self, node: BagNode, parent: Any) -> Any:
        """Compile image node. Parent is the Document."""
        attrs = node.runtime_attrs
        path = attrs.get("path", "")
        width = attrs.get("width")
        height = attrs.get("height")
        align = attrs.get("align")

        if not path:
            return None

        para = parent.add_paragraph()
        if align:
            para.alignment = self._get_alignment(align)

        docx_run = para.add_run()
        if width and height:
            docx_run.add_picture(str(path), width=Inches(width), height=Inches(height))
        elif width:
            docx_run.add_picture(str(path), width=Inches(width))
        elif height:
            docx_run.add_picture(str(path), height=Inches(height))
        else:
            docx_run.add_picture(str(path))

        return para

    @compiler()
    def pagebreak(self, node: BagNode, parent: Any) -> Any:  # noqa: ARG002
        """Compile pagebreak node."""
        parent.add_page_break()
        return None

    @compiler()
    def header(self, node: BagNode, parent: Any) -> Any:  # noqa: ARG002
        """Compile header node. Parent is the Document."""
        section = parent.sections[0]
        hdr = section.header
        hdr.is_linked_to_previous = False
        return hdr

    @compiler()
    def footer(self, node: BagNode, parent: Any) -> Any:  # noqa: ARG002
        """Compile footer node. Parent is the Document."""
        section = parent.sections[0]
        ftr = section.footer
        ftr.is_linked_to_previous = False
        return ftr

    # -------------------------------------------------------------------------
    # Formatting helpers
    # -------------------------------------------------------------------------

    def _apply_run_formatting(self, attrs: dict[str, Any], run: Any) -> None:
        """Apply formatting attributes to a run."""
        bold = attrs.get("bold", False)
        if bold:
            run.bold = True

        italic = attrs.get("italic", False)
        if italic:
            run.italic = True

        underline = attrs.get("underline", False)
        if underline:
            run.underline = True

        strike = attrs.get("strike", False)
        if strike:
            run.font.strike = True

        font_size = attrs.get("font_size")
        if font_size:
            run.font.size = Pt(font_size)

        font_name = attrs.get("font_name")
        if font_name:
            run.font.name = str(font_name)

        color = attrs.get("color")
        if color:
            run.font.color.rgb = RGBColor.from_string(str(color))

        highlight = attrs.get("highlight")
        if highlight:
            highlight_map = {
                "yellow": WD_COLOR_INDEX.YELLOW,
                "green": WD_COLOR_INDEX.BRIGHT_GREEN,
                "cyan": WD_COLOR_INDEX.TURQUOISE,
                "magenta": WD_COLOR_INDEX.PINK,
                "blue": WD_COLOR_INDEX.BLUE,
                "red": WD_COLOR_INDEX.RED,
                "gray": WD_COLOR_INDEX.GRAY_25,
            }
            if str(highlight).lower() in highlight_map:
                run.font.highlight_color = highlight_map[str(highlight).lower()]

    def _get_alignment(self, align: str) -> Any:
        """Convert alignment string to WD_ALIGN_PARAGRAPH."""
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return align_map.get(str(align).lower(), WD_ALIGN_PARAGRAPH.LEFT)

    def _set_cell_shading(self, cell: Any, color: str) -> None:
        """Set cell background color."""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shading)


# Set _compiler_class after WordCompiler is defined
from genro_office.builders.word_builder import WordBuilder  # noqa: E402

WordBuilder._compiler_class = WordCompiler
